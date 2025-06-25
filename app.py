"""
RFP Analyzer AI Agent - Streamlit App

Deployment:
1. Set your Gemini API key as an environment variable: GEMINI_API_KEY
   (e.g., export GEMINI_API_KEY=your-key on Linux/Mac, set GEMINI_API_KEY=your-key on Windows)
2. Install dependencies: pip install -r requirements.txt
3. Run: streamlit run app.py
"""

import os
import streamlit as st
import json
import tempfile
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import google.generativeai as genai
from datetime import datetime
from typing import Dict, Any
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# --- Gemini API Key from environment ---
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    st.error("GEMINI_API_KEY environment variable not set. Please set it before running the app.")
    st.stop()

# --- RFPAnalyzer class ---
class RFPAnalyzer:
    def __init__(self, api_key: str):
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-2.0-flash')
        self.generation_config = genai.types.GenerationConfig(
            temperature=0.1,
            max_output_tokens=8192,
        )

    def extract_text_from_pdf(self, pdf_file) -> str:
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(pdf_file.read())
                tmp_path = tmp_file.name
            doc = fitz.open(tmp_path)
            text = ""
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                if len(page_text.strip()) < 100:
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as img_file:
                        img_file.write(img_data)
                        img_path = img_file.name
                    ocr_text = pytesseract.image_to_string(Image.open(img_path))
                    text += f"\n--- Page {page_num + 1} (OCR) ---\n{ocr_text}\n"
                    os.unlink(img_path)
                else:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            doc.close()
            os.unlink(tmp_path)
            return text
        except Exception as e:
            st.error(f"Error extracting text from PDF: {str(e)}")
            return ""

    def extract_text_from_image(self, image_file) -> str:
        try:
            image = Image.open(image_file)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            st.error(f"Error extracting text from image: {str(e)}")
            return ""

    def analyze_rfp(self, text: str) -> Dict[str, Any]:
        """Analyze RFP text and extract structured information"""

        analysis_prompt = f"""You are an expert RFP analyst for contractor companies.
Analyze the provided RFP document and extract comprehensive insights in a structured JSON format.
Focus on information that helps contractors understand requirements and prepare competitive bids.

Analyze this RFP document and provide detailed insights in the following JSON structure:

{{
    "project_overview": {{
        "title": "Project title",
        "description": "Brief project description",
        "client_organization": "Client name/organization",
        "project_type": "Type of project",
        "industry_sector": "Industry/sector"
    }},
    "requirements": {{
        "functional_requirements": ["List of functional requirements"],
        "technical_requirements": ["List of technical requirements"],
        "compliance_requirements": ["Regulatory/compliance needs"],
        "performance_requirements": ["Performance criteria"]
    }},
    "technology_stack": {{
        "preferred_technologies": ["Preferred tech stack"],
        "platforms": ["Required platforms"],
        "databases": ["Database requirements"],
        "frameworks": ["Framework preferences"],
        "third_party_integrations": ["Required integrations"]
    }},
    "project_details": {{
        "estimated_budget": "Budget range if mentioned",
        "timeline": "Project timeline",
        "start_date": "Expected start date",
        "key_milestones": ["Important milestones"],
        "deliverables": ["Expected deliverables"]
    }},
    "project_phases": {{
        "suggested_phases": ["Recommended development phases"],
        "phase_descriptions": ["Description of each phase"]
    }},
    "evaluation_criteria": {{
        "technical_criteria": ["Technical evaluation factors"],
        "commercial_criteria": ["Cost evaluation factors"],
        "experience_criteria": ["Experience requirements"],
        "weightage": "Scoring weightage if mentioned"
    }},
    "submission_requirements": {{
        "proposal_format": "Required proposal format",
        "submission_deadline": "Deadline for submission",
        "required_documents": ["Required documents"],
        "contact_information": "Contact details"
    }},
    "risk_analysis": {{
        "technical_risks": ["Potential technical challenges"],
        "project_risks": ["Project delivery risks"],
        "mitigation_strategies": ["Suggested risk mitigation"]
    }},
    "competitive_analysis": {{
        "likely_competitors": ["Potential competing companies"],
        "competitive_advantages": ["Areas to highlight"],
        "differentiators": ["Unique selling points to emphasize"]
    }},
    "bid_strategy": {{
        "key_strengths_to_highlight": ["Strengths to emphasize"],
        "pricing_strategy": "Recommended pricing approach",
        "proposal_focus_areas": ["Areas to focus on in proposal"],
        "win_probability": "Estimated win probability and reasoning"
    }}
}}

RFP Document:
{text}

Provide only the JSON response with detailed analysis. Be thorough and extract as much relevant information as possible.
"""

        try:
            response = self.model.generate_content(
                analysis_prompt,
                generation_config=self.generation_config
            )

            # Clean and parse JSON response
            content = response.text.strip()

            # Remove markdown code blocks if present
            if content.startswith('```json'):
                content = content[7:]
            if content.endswith('```'):
                content = content[:-3]

            # Clean any remaining markdown
            content = content.strip()

            analysis = json.loads(content)
            return analysis

        except json.JSONDecodeError as e:
            st.error(f"Error parsing AI response: {str(e)}")
            # Try to extract JSON from the response if it's embedded
            try:
                # Look for JSON pattern in the response
                json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
                if json_match:
                    analysis = json.loads(json_match.group())
                    return analysis
            except:
                pass
            return {"error": "Failed to parse analysis"}
        except Exception as e:
            st.error(f"Error analyzing RFP: {str(e)}")
            return {"error": str(e)}

    def generate_executive_summary(self, analysis: Dict[str, Any]) -> str:
        """Generate an executive summary of the RFP analysis"""

        summary_prompt = f"""You are an expert business analyst. Create a concise executive summary for contractors.

Based on this RFP analysis, create a brief executive summary (200-300 words) that highlights:
1. Key project opportunity
2. Critical requirements
3. Main challenges and risks
4. Recommended bid approach
5. Win probability assessment

Analysis data:
{json.dumps(analysis, indent=2)}

Write in a professional, actionable tone for decision-makers.
"""

        try:
            response = self.model.generate_content(
                summary_prompt,
                generation_config=self.generation_config
            )
            return response.text
        except Exception as e:
            return f"Error generating summary: {str(e)}"

    def generate_docx_report(self, analysis: Dict[str, Any], summary: str) -> BytesIO:
        """Generate a comprehensive DOCX report"""
        try:
            # Create a new Document
            doc = Document()

            # Add title
            title = doc.add_heading('RFP Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add timestamp
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")

            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(summary)
            doc.add_page_break()

            # Project Overview
            if 'project_overview' in analysis:
                doc.add_heading('Project Overview', level=1)
                overview = analysis['project_overview']

                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Attribute'
                hdr_cells[1].text = 'Details'

                for key, value in overview.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = key.replace('_', ' ').title()
                    row_cells[1].text = str(value) if value else 'N/A'

                doc.add_paragraph("")

            # Requirements
            if 'requirements' in analysis:
                doc.add_heading('Requirements', level=1)
                requirements = analysis['requirements']

                for req_type, req_list in requirements.items():
                    if req_list:
                        doc.add_heading(req_type.replace('_', ' ').title(), level=2)
                        for req in req_list:
                            doc.add_paragraph(f"• {req}")
                        doc.add_paragraph("")

            # Technology Stack
            if 'technology_stack' in analysis:
                doc.add_heading('Technology Stack', level=1)
                tech_stack = analysis['technology_stack']

                for tech_type, tech_list in tech_stack.items():
                    if tech_list:
                        doc.add_heading(tech_type.replace('_', ' ').title(), level=2)
                        for tech in tech_list:
                            doc.add_paragraph(f"• {tech}")
                        doc.add_paragraph("")

            # Project Details
            if 'project_details' in analysis:
                doc.add_heading('Project Details', level=1)
                details = analysis['project_details']

                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Attribute'
                hdr_cells[1].text = 'Details'

                for key, value in details.items():
                    if key == 'key_milestones' or key == 'deliverables':
                        if value:
                            row_cells = table.add_row().cells
                            row_cells[0].text = key.replace('_', ' ').title()
                            row_cells[1].text = '\n'.join([f"• {item}" for item in value])
                    else:
                        row_cells = table.add_row().cells
                        row_cells[0].text = key.replace('_', ' ').title()
                        row_cells[1].text = str(value) if value else 'N/A'

                doc.add_paragraph("")

            # Bid Strategy
            if 'bid_strategy' in analysis:
                doc.add_heading('Bid Strategy Recommendations', level=1)
                strategy = analysis['bid_strategy']

                for key, value in strategy.items():
                    if value:
                        doc.add_heading(key.replace('_', ' ').title(), level=2)
                        if isinstance(value, list):
                            for item in value:
                                doc.add_paragraph(f"• {item}")
                        else:
                            doc.add_paragraph(str(value))
                        doc.add_paragraph("")

            # Risk Analysis
            if 'risk_analysis' in analysis:
                doc.add_heading('Risk Analysis', level=1)
                risks = analysis['risk_analysis']

                for risk_type, risk_list in risks.items():
                    if risk_list:
                        doc.add_heading(risk_type.replace('_', ' ').title(), level=2)
                        if isinstance(risk_list, list):
                            for risk in risk_list:
                                doc.add_paragraph(f"• {risk}")
                        else:
                            doc.add_paragraph(str(risk_list))
                        doc.add_paragraph("")

            # Submission Requirements
            if 'submission_requirements' in analysis:
                doc.add_heading('Submission Requirements', level=1)
                submission = analysis['submission_requirements']

                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Requirement'
                hdr_cells[1].text = 'Details'

                for key, value in submission.items():
                    if key == 'required_documents' and value:
                        row_cells = table.add_row().cells
                        row_cells[0].text = key.replace('_', ' ').title()
                        row_cells[1].text = '\n'.join([f"• {doc}" for doc in value])
                    else:
                        row_cells = table.add_row().cells
                        row_cells[0].text = key.replace('_', ' ').title()
                        row_cells[1].text = str(value) if value else 'N/A'

            # Save to BytesIO
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)

            return docx_buffer

        except Exception as e:
            st.error(f"Error generating DOCX report: {str(e)}")
            return None

# --- Streamlit UI ---
st.set_page_config(
    page_title="RFP Analyzer AI Agent",
    page_icon="📄",
    layout="wide"
)
st.title("🤖 RFP Analysis AI Agent")
st.markdown("*Intelligent RFP analysis for contractor companies*")

with st.sidebar:
    st.header("Configuration")
    st.markdown(f"**Gemini API Key:** {'Set' if GEMINI_API_KEY else 'Not Set'}")
    st.markdown("---")
    st.markdown("**Supported Formats:**")
    st.markdown("• PDF documents")
    st.markdown("• Images (PNG, JPG, JPEG)")
    st.markdown("---")
    st.markdown("**AI Model:**")
    st.markdown("• Gemini 2.0 Flash")
    if st.button("Clear Analysis", type="secondary"):
        for k in ["analysis_result", "extracted_text", "executive_summary"]:
            if k in st.session_state:
                del st.session_state[k]
        st.rerun()

analyzer = RFPAnalyzer(GEMINI_API_KEY)

st.header("📤 Upload RFP Document")
uploaded_file = st.file_uploader(
    "Choose an RFP document",
    type=['pdf', 'png', 'jpg', 'jpeg'],
    help="Upload a PDF or image file containing the RFP"
)

if uploaded_file is not None:
    st.success(f"File uploaded: {uploaded_file.name} ({uploaded_file.size} bytes)")
    if st.button("🔍 Analyze RFP", type="primary"):
        with st.spinner("Extracting text from document..."):
            if uploaded_file.type == "application/pdf":
                extracted_text = analyzer.extract_text_from_pdf(uploaded_file)
            else:
                extracted_text = analyzer.extract_text_from_image(uploaded_file)
            st.session_state.extracted_text = extracted_text
        if extracted_text:
            with st.spinner("Analyzing RFP with AI... This may take a few minutes."):
                analysis = analyzer.analyze_rfp(extracted_text)
                st.session_state.analysis_result = analysis
            st.success("✅ Analysis completed!")

if "analysis_result" in st.session_state and "extracted_text" in st.session_state:
    analysis = st.session_state.analysis_result
    if "error" not in analysis:
        st.header("📋 Executive Summary")
        with st.spinner("Generating executive summary..."):
            if "executive_summary" not in st.session_state:
                summary = analyzer.generate_executive_summary(analysis)
                st.session_state.executive_summary = summary
            else:
                summary = st.session_state.executive_summary
            st.markdown(summary)
        st.markdown("---")
        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "📊 Project Overview",
            "⚙️ Technical Details",
            "💼 Business Details",
            "🎯 Bid Strategy",
            "📄 Raw Text"
        ])
        with tab1:
            st.subheader("Project Overview")
            if "project_overview" in analysis:
                overview = analysis["project_overview"]
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**Title:** {overview.get('title', 'N/A')}")
                    st.write(f"**Client:** {overview.get('client_organization', 'N/A')}")
                    st.write(f"**Type:** {overview.get('project_type', 'N/A')}")
                with col2:
                    st.write(f"**Industry:** {overview.get('industry_sector', 'N/A')}")
                    st.write(f"**Description:** {overview.get('description', 'N/A')}")
            st.subheader("Requirements")
            if "requirements" in analysis:
                req = analysis["requirements"]
                if req.get("functional_requirements"):
                    st.write("**Functional Requirements:**")
                    for item in req["functional_requirements"]:
                        st.write(f"• {item}")
                if req.get("technical_requirements"):
                    st.write("**Technical Requirements:**")
                    for item in req["technical_requirements"]:
                        st.write(f"• {item}")
        with tab2:
            st.subheader("Technology Stack")
            if "technology_stack" in analysis:
                tech = analysis["technology_stack"]
                col1, col2 = st.columns(2)
                with col1:
                    if tech.get("preferred_technologies"):
                        st.write("**Preferred Technologies:**")
                        for item in tech["preferred_technologies"]:
                            st.write(f"• {item}")
                    if tech.get("platforms"):
                        st.write("**Platforms:**")
                        for item in tech["platforms"]:
                            st.write(f"• {item}")
                with col2:
                    if tech.get("databases"):
                        st.write("**Databases:**")
                        for item in tech["databases"]:
                            st.write(f"• {item}")
                    if tech.get("third_party_integrations"):
                        st.write("**Integrations:**")
                        for item in tech["third_party_integrations"]:
                            st.write(f"• {item}")
            st.subheader("Project Phases")
            if "project_phases" in analysis:
                phases = analysis["project_phases"]
                if phases.get("suggested_phases"):
                    for i, phase in enumerate(phases["suggested_phases"], 1):
                        st.write(f"**Phase {i}:** {phase}")
        with tab3:
            st.subheader("Project Details")
            if "project_details" in analysis:
                details = analysis["project_details"]
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**Budget:** {details.get('estimated_budget', 'N/A')}")
                    st.write(f"**Timeline:** {details.get('timeline', 'N/A')}")
                    st.write(f"**Start Date:** {details.get('start_date', 'N/A')}")
                with col2:
                    if details.get("key_milestones"):
                        st.write("**Key Milestones:**")
                        for milestone in details["key_milestones"]:
                            st.write(f"• {milestone}")
            st.subheader("Submission Requirements")
            if "submission_requirements" in analysis:
                sub = analysis["submission_requirements"]
                st.write(f"**Deadline:** {sub.get('submission_deadline', 'N/A')}")
                st.write(f"**Contact:** {sub.get('contact_information', 'N/A')}")
                if sub.get("required_documents"):
                    st.write("**Required Documents:**")
                    for doc in sub["required_documents"]:
                        st.write(f"• {doc}")
        with tab4:
            st.subheader("Bid Strategy Recommendations")
            if "bid_strategy" in analysis:
                strategy = analysis["bid_strategy"]
                col1, col2 = st.columns(2)
                with col1:
                    if strategy.get("key_strengths_to_highlight"):
                        st.write("**Strengths to Highlight:**")
                        for strength in strategy["key_strengths_to_highlight"]:
                            st.write(f"• {strength}")
                    st.write(f"**Pricing Strategy:** {strategy.get('pricing_strategy', 'N/A')}")
                with col2:
                    if strategy.get("proposal_focus_areas"):
                        st.write("**Proposal Focus Areas:**")
                        for area in strategy["proposal_focus_areas"]:
                            st.write(f"• {area}")
                    st.write(f"**Win Probability:** {strategy.get('win_probability', 'N/A')}")
            st.subheader("Risk Analysis")
            if "risk_analysis" in analysis:
                risks = analysis["risk_analysis"]
                if risks.get("technical_risks"):
                    st.write("**Technical Risks:**")
                    for risk in risks["technical_risks"]:
                        st.write(f"• {risk}")
                if risks.get("mitigation_strategies"):
                    st.write("**Mitigation Strategies:**")
                    for strategy in risks["mitigation_strategies"]:
                        st.write(f"• {strategy}")
        with tab5:
            st.subheader("Extracted Text")
            st.text_area("Raw extracted text:",
                       value=st.session_state.extracted_text,
                       height=400,
                       disabled=True)
            st.subheader("Download Results")
            col1, col2, col3 = st.columns(3)
            with col1:
                json_str = json.dumps(analysis, indent=2)
                st.download_button(
                    label="📥 Download Analysis (JSON)",
                    data=json_str,
                    file_name=f"rfp_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json"
                )
            with col2:
                st.download_button(
                    label="📥 Download Extracted Text",
                    data=st.session_state.extracted_text,
                    file_name=f"extracted_text_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain"
                )
            with col3:
                if st.button("📥 Generate DOCX Report"):
                    with st.spinner("Generating DOCX report..."):
                        docx_buffer = analyzer.generate_docx_report(analysis, st.session_state.executive_summary)
                        if docx_buffer:
                            st.download_button(
                                label="📥 Download DOCX Report",
                                data=docx_buffer.getvalue(),
                                file_name=f"rfp_analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        else:
                            st.error("Failed to generate DOCX report")
    else:
        st.error(f"Analysis failed: {analysis.get('error', 'Unknown error')}")
st.markdown("---")
st.markdown("*Built with Google Gemini • Made for Contractor Companies*")
